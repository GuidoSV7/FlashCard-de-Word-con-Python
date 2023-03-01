from docx import Document
from docx.shared import Cm
from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from docx.enum.table import WD_ALIGN_VERTICAL




vector_palabras = ['¿Qués es VLSM?',
                   'VLSM representan otra de las tantas soluciones que se implementaron para evitar el agotamiento de direcciones IP en IPv4', 
                   'La principal característica de 5G es proveer banda ancha a alta velocidad',
                   'Verdadero', 
                   '¿Qué significa UTP?',
                   'Cable de par trenzado sin blindaje']

def create_rectangles(document, palabras, width, height):
    # Creamos la carpeta para las imágenes si no existe
    if not os.path.exists('Imagenes'):
        os.mkdir('Imagenes')

    # Calculamos el ancho de cada celda de la tabla
    num_rectangles = len(palabras)
    cell_width = Cm(7 / num_rectangles)

    # Creamos una tabla con 2 columnas y suficientes filas para acomodar los rectángulos
    num_rows = (num_rectangles + 1) // 2
    table = document.add_table(rows=num_rows, cols=2)

    # Centramos la tabla
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregamos una imagen en cada celda de la tabla
    for i, palabra in enumerate(palabras):
        # Creamos una imagen en blanco
        image = Image.new('RGB', (width, height), color='white')

        # agregar borde negro de 5 píxeles
        bordered_image = ImageOps.expand(image, border=5, fill="black")

        # Dibujamos un rectángulo en la imagen
        draw = ImageDraw.Draw(bordered_image)
        rectangle = [(105, 105), (595, 295)]  # cambiamos las coordenadas para agregar el borde
        draw.rectangle(rectangle)

        # Agregamos texto dentro del rectángulo
        font_size = 50
        font = ImageFont.truetype('arial.ttf', size=font_size) # especificamos el tamaño de fuente
        text_width, text_height = draw.textsize(palabra, font=font) # obtenemos las dimensiones del texto
        while text_width > width - 20 or text_height > height - 20: # ajustamos tamaño de fuente si no cabe en el rectángulo
            font_size -= 1
            font = ImageFont.truetype('arial.ttf', size=font_size)
            text_width, text_height = draw.textsize(palabra, font=font)
        x = (width - text_width) / 2 # calculamos la coordenada X para centrar el texto
        y = (height - text_height) / 2 # calculamos la coordenada Y para centrar el texto
        draw.text((x, y), palabra, fill='black', font=font) # agregamos el parámetro font

        # Guardamos la imagen en la carpeta "Imagenes"
        image_filename = f'Imagenes/rectangulo_{i+1}.png'
        bordered_image.save(image_filename)

        # Agregamos la imagen en la celda correspondiente de la tabla
        cell = table.cell(i // 2, i % 2)
        cell_width = Cm(7)
        cell.width = cell_width
        cell.height = Cm(3)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_filename, width=cell_width)

    # Agregamos un salto de página después de la tabla
    document.add_page_break()

# Creamos el documento
document = Document()

# Definimos las dimensiones de los rectángulos
width, height = 700, 300


# Crear rectángulos
create_rectangles(document, vector_palabras, width, height)


# Guardamos el documento
document.save('rectangulos.docx')